
import os
import sys
import shutil
import zipfile
import traceback
import sqlite3
from datetime import datetime
from typing import List, Tuple, Dict, Any

import pandas as pd
from openpyxl import load_workbook
from docx import Document

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
    QLineEdit, QPushButton, QTableWidget, QTableWidgetItem, QFileDialog,
    QMessageBox, QComboBox, QInputDialog, QTreeWidget, QTreeWidgetItem,
    QTabWidget, QListWidget, QProgressDialog, QSizePolicy, QSplitter
)
from PyQt6.QtGui import QPixmap, QFont, QMovie
from PyQt6.QtCore import Qt

# Optional dark style availability (not required)
try:
    import qdarkstyle
    QDARK_AVAILABLE = True
except Exception:
    QDARK_AVAILABLE = False

DB_FILE = "npi_projects.db"
EXCEL_FILE = "NPI_Project_Data.xlsx"
PRODUCT_SHEET = "Products"
TABLE_ROWS = 9
ASSEMBLY_ROWS = 9
MACHINE_ROWS = 9

PROFESSIONAL_QSS = """ /* same QSS as before - truncated for brevity in code block */
QWidget { font-family: "Segoe UI", Arial, sans-serif; font-size: 11px; background-color: #181b1f; color: #f1f3f5; }
QLabel#HeaderLabel { font-size: 15px; font-weight: 700; color: #ffffff; }
QPushButton { border: 1px solid #3a454d; background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #2a3238, stop:1 #1f2428); color: #f1f3f5; padding: 6px 14px; border-radius: 8px; font-weight: 500; }
QPushButton:hover { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #334148, stop:1 #252c31); border: 1px solid #4c5a64; }
QPushButton:pressed { background: #2f6db0; border: 1px solid #2f6db0; color: #ffffff; }
QLineEdit, QComboBox, QListWidget, QTableWidget, QTreeWidget { border: 1px solid #3a454d; border-radius: 6px; padding: 6px; background: #202428; color: #f8f9fa; selection-background-color: #2f6db0; selection-color: #ffffff; }
QHeaderView::section { background: #2a3238; color: #e6eef3; padding: 6px; border: 1px solid #3a454d; border-radius: 6px; font-weight: 500; }
QTabWidget::pane { border-top: 1px solid #3a454d; background: #1f2327; border-radius: 6px; }
QTabBar::tab { background: #1c1f24; padding: 8px 16px; margin-right: 4px; border-top-left-radius: 8px; border-top-right-radius: 8px; color: #cfd8dc; }
QTabBar::tab:hover { background: #2a3238; color: #ffffff; }
QTabBar::tab:selected { background: #2f6db0; border: 1px solid #2f6db0; border-bottom: 0px; color: #ffffff; font-weight: 600; }
"""

# ----------------- Database Manager -----------------
class DBManager:
    def __init__(self, db_file: str = DB_FILE):
        self.db_file = db_file
        self.conn = sqlite3.connect(self.db_file, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        self._create_tables()

    def _create_tables(self):
        sql = """
        CREATE TABLE IF NOT EXISTS projects (
            project_id INTEGER PRIMARY KEY AUTOINCREMENT,
            product_name TEXT,
            project_name TEXT UNIQUE,
            fg_part_number TEXT,
            pcba_part_number TEXT,
            start_date TEXT,
            end_date TEXT,
            bom_file TEXT,
            npi_engineer TEXT
        );
        CREATE TABLE IF NOT EXISTS mes (
            mes_id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            lot_id TEXT,
            workflow_smt TEXT,
            workflow_tla TEXT,
            smt_work_order TEXT,
            tla_work_order TEXT,
            work_order_qty INTEGER,
            po_number TEXT,
            po_qty INTEGER,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS assembly_drawings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            assembly_drawing TEXT,
            drawing_name TEXT,
            seq INTEGER,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS build_matrix (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            component TEXT,
            make TEXT,
            seq INTEGER,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS machine_matrix (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            machine_name TEXT,
            program_name TEXT,
            seq INTEGER,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS handover_docs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            category TEXT,
            file_path TEXT,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS checklist_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER,
            item_name TEXT,
            completed INTEGER DEFAULT 0,
            person TEXT,
            reference TEXT,
            seq INTEGER,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        """
        cur = self.conn.cursor()
        cur.executescript(sql)
        self.conn.commit()

    # Projects
    def add_project(self, product_name: str, project_name: str, details: Dict[str, Any]) -> int:
        cur = self.conn.cursor()

        # Insert main project
        cur.execute("""
            INSERT OR IGNORE INTO projects (product_name, project_name, fg_part_number, pcba_part_number, start_date, end_date, bom_file, npi_engineer)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            product_name, project_name,
            details.get("FG Part Number", ""),
            details.get("PCBA Part Number", ""),
            details.get("Start Date", ""),
            details.get("End Date", ""),
            details.get("BOM File", ""),
            details.get("NPI Engineer", "")
        ))
        self.conn.commit()

        # Get project_id
        cur.execute("SELECT project_id FROM projects WHERE project_name=?", (project_name,))
        row = cur.fetchone()
        if not row:
            return -1
        project_id = row["project_id"]

        # Insert MES Workflow
        workflow = details.get("MES Workflow", {})
        if workflow:
            cur.execute("""
                INSERT INTO mes_workflow (project_id, lot_id, workflow_smt, workflow_tla, smt_work_order, tla_work_order, work_order_qty, po_number, po_qty)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                project_id,
                workflow.get("LOT ID", ""),
                workflow.get("Workflow SMT - Name", ""),
                workflow.get("Workflow TLA - Name", ""),
                workflow.get("SMT - Work Order", ""),
                workflow.get("TLA - Work Order", ""),
                workflow.get("Work Order Quantity", 0),
                workflow.get("PO NUMBER", ""),
                workflow.get("PO Quantity", 0),
            ))

        # Insert Assembly Drawings
        for drawing in details.get("Assembly Drawings", []):
            cur.execute("""
                INSERT INTO assembly_drawings (project_id, assembly_drawing, drawing_name)
                VALUES (?, ?, ?)
            """, (project_id, drawing.get("Assembly Drawing", ""), drawing.get("Drawing Name", "")))

        # Insert Build Matrix
        for item in details.get("Build Matrix", []):
            cur.execute("""
                INSERT INTO build_matrix (project_id, component, make)
                VALUES (?, ?, ?)
            """, (project_id, item.get("Component", ""), item.get("Make", "")))

        # Insert Machine Program
        for prog in details.get("Machine Program", []):
            cur.execute("""
                INSERT INTO machine_program (project_id, machine_name, program_name)
                VALUES (?, ?, ?)
            """, (project_id, prog.get("Machine Name", ""), prog.get("Program Name", "")))

        self.conn.commit()
        return project_id


    def update_project_details(self, project_name: str, details: Dict[str, Any]):
        cur = self.conn.cursor()
        cur.execute("""
            UPDATE projects SET fg_part_number=?, pcba_part_number=?, start_date=?, end_date=?, bom_file=?, npi_engineer=?
            WHERE project_name=?
        """, (details.get("FG Part Number", ""), details.get("PCBA Part Number", ""),
              details.get("Start Date", ""), details.get("End Date", ""), details.get("BOM File", ""), details.get("NPI Engineer", ""), project_name))
        self.conn.commit()

    def list_products(self) -> List[str]:
        cur = self.conn.cursor()
        cur.execute("SELECT DISTINCT product_name FROM projects WHERE product_name IS NOT NULL AND product_name!='' ORDER BY product_name")
        return [r["product_name"] for r in cur.fetchall()]

    def list_projects_for_product(self, product_name: str) -> List[str]:
        cur = self.conn.cursor()
        cur.execute("SELECT project_name FROM projects WHERE product_name=? ORDER BY project_name", (product_name,))
        return [r["project_name"] for r in cur.fetchall()]

    def get_project_by_name(self, project_name: str) -> sqlite3.Row:
        cur = self.conn.cursor()
        cur.execute("SELECT * FROM projects WHERE project_name=?", (project_name,))
        return cur.fetchone()

    def delete_project(self, project_name: str):
        cur = self.conn.cursor()
        cur.execute("SELECT project_id FROM projects WHERE project_name=?", (project_name,))
        row = cur.fetchone()
        if row:
            pid = row["project_id"]
            cur.execute("DELETE FROM projects WHERE project_id=?", (pid,))
            cur.execute("DELETE FROM mes WHERE project_id=?", (pid,))
            cur.execute("DELETE FROM build_matrix WHERE project_id=?", (pid,))
            cur.execute("DELETE FROM machine_matrix WHERE project_id=?", (pid,))
            cur.execute("DELETE FROM handover_docs WHERE project_id=?", (pid,))
            cur.execute("DELETE FROM checklist_items WHERE project_id=?", (pid,))
            self.conn.commit()

    # MES
    def save_mes(self, project_id: int, mes: Dict[str, Any]):
        cur = self.conn.cursor()
        cur.execute("DELETE FROM mes WHERE project_id=?", (project_id,))
        cur.execute("""
            INSERT INTO mes (project_id, lot_id, workflow_smt, workflow_tla, smt_work_order, tla_work_order, work_order_qty, po_number, po_qty)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (project_id, mes.get("LOT ID", ""), mes.get("Workflow SMT - Name", ""), mes.get("Workflow TLA - Name", ""),
              mes.get("SMT - Work Order", ""), mes.get("TLA - Work Order", ""), mes.get("Work Order Quantity", None),
              mes.get("PO NUMBER", ""), mes.get("PO Quantity", None)))
        self.conn.commit()

    def get_mes(self, project_id: int) -> sqlite3.Row:
        cur = self.conn.cursor()
        cur.execute("SELECT * FROM mes WHERE project_id=?", (project_id,))
        return cur.fetchone()

    # Build matrix
    def save_build_matrix(self, project_id: int, rows: List[Tuple[str, str]]):
        cur = self.conn.cursor()
        cur.execute("DELETE FROM build_matrix WHERE project_id=?", (project_id,))
        for seq, (comp, make) in enumerate(rows, start=1):
            cur.execute("INSERT INTO build_matrix (project_id, component, make, seq) VALUES (?, ?, ?, ?)",
                        (project_id, comp or "", make or "", seq))
        self.conn.commit()

    def get_build_matrix(self, project_id: int) -> List[sqlite3.Row]:
        cur = self.conn.cursor()
        cur.execute("SELECT component, make FROM build_matrix WHERE project_id=? ORDER BY seq", (project_id,))
        return cur.fetchall()
    # Assembly matrix
    def save_assembly_drawings(self, project_id: int, rows: list[tuple[str, str]]):
        cur = self.conn.cursor()
        # Delete existing rows for this project
        cur.execute("DELETE FROM assembly_drawings WHERE project_id=?", (project_id,))
        # Insert new rows
        for seq, (assembly_drawing, drawing_name) in enumerate(rows, start=1):
            cur.execute(
                "INSERT INTO assembly_drawings (project_id, assembly_drawing, drawing_name, seq) VALUES (?, ?, ?, ?)",
                (project_id, assembly_drawing or "", drawing_name or "", seq)
            )
        self.conn.commit()

    def get_assembly_drawings(self, project_id: int) -> list[sqlite3.Row]:
        cur = self.conn.cursor()
        cur.execute(
            "SELECT assembly_drawing, drawing_name FROM assembly_drawings WHERE project_id=? ORDER BY seq",
            (project_id,)
        )
        return cur.fetchall()

    # Machine matrix
    def save_machine_matrix(self, project_id: int, rows: List[Tuple[str, str]]):
        cur = self.conn.cursor()
        cur.execute("DELETE FROM machine_matrix WHERE project_id=?", (project_id,))
        for seq, (mn, pn) in enumerate(rows, start=1):
            cur.execute("INSERT INTO machine_matrix (project_id, machine_name, program_name, seq) VALUES (?, ?, ?, ?)",
                        (project_id, mn or "", pn or "", seq))
        self.conn.commit()

    def get_machine_matrix(self, project_id: int) -> List[sqlite3.Row]:
        cur = self.conn.cursor()
        cur.execute("SELECT machine_name, program_name FROM machine_matrix WHERE project_id=? ORDER BY seq", (project_id,))
        return cur.fetchall()

    # Handover docs
    def add_handover_doc(self, project_id: int, category: str, file_path: str):
        cur = self.conn.cursor()
        cur.execute("INSERT INTO handover_docs (project_id, category, file_path) VALUES (?, ?, ?)", (project_id, category, file_path))
        self.conn.commit()

    def get_handover_docs(self, project_id: int, category: str = None) -> List[sqlite3.Row]:
        cur = self.conn.cursor()
        if category:
            cur.execute("SELECT * FROM handover_docs WHERE project_id=? AND category=? ORDER BY id", (project_id, category))
        else:
            cur.execute("SELECT * FROM handover_docs WHERE project_id=? ORDER BY category, id", (project_id,))
        return cur.fetchall()

    def remove_handover_doc(self, doc_id: int):
        cur = self.conn.cursor()
        cur.execute("DELETE FROM handover_docs WHERE id=?", (doc_id,))
        self.conn.commit()

    # Checklist
    def initialize_checklist(self, project_id: int, template: Dict[str, Dict[str, Any]]):
        # If no checklist exists, populate using template
        cur = self.conn.cursor()
        cur.execute("SELECT COUNT(*) as c FROM checklist_items WHERE project_id=?", (project_id,))
        if cur.fetchone()["c"] == 0:
            seq = 1
            for name, entry in template.items():
                cur.execute("INSERT INTO checklist_items (project_id, item_name, completed, person, reference, seq) VALUES (?, ?, ?, ?, ?, ?)",
                            (project_id, name, int(entry.get("completed", False)), entry.get("person", ""), entry.get("reference", ""), seq))
                seq += 1
            self.conn.commit()

    def get_checklist(self, project_id: int) -> List[sqlite3.Row]:
        cur = self.conn.cursor()
        cur.execute("SELECT * FROM checklist_items WHERE project_id=? ORDER BY seq", (project_id,))
        return cur.fetchall()

    def update_checklist_item(self, item_id: int, completed: int, person: str, reference: str):
        cur = self.conn.cursor()
        cur.execute("UPDATE checklist_items SET completed=?, person=?, reference=? WHERE id=?", (completed, person, reference, item_id))
        self.conn.commit()

    # Utilities
    def close(self):
        self.conn.close()


# ----------------- Handover Tab -----------------
class HandoverTab(QWidget):
    def __init__(self, db: DBManager, get_project_dir_callable, get_project_info_callable, logo_path="tsat.png", parent=None):
        super().__init__(parent)
        self.db = db
        self.get_project_dir = get_project_dir_callable
        self.get_project_info = get_project_info_callable
        self.logo_path = logo_path
        self.lists_widgets: Dict[str, QListWidget] = {}
        self.categories = [
            "Process Flow Chart","PFMEA", "Control Plan", "Process Parameters", "SAP BOM", "Label Artwork",
            "Cycle Time Study", "Assembly Qualification Report", "Packaging Document",
            "WI", "SOP", "Stencil, Tools & Fixtures", "Lessons Learnt", "Other Documents"
        ]
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout()
        header = QLabel("Handover Documents")
        header.setObjectName("HeaderLabel")
        layout.addWidget(header)

        for cat in self.categories:
            row = QHBoxLayout()
            label = QLabel(cat)
            label.setFixedWidth(180)
            label.setFont(QFont("Segoe UI", 10, QFont.Weight.DemiBold))
            lw = QListWidget()
            lw.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)

            btns = QVBoxLayout()
            btn_add = QPushButton("Add Files")
            btn_add.clicked.connect(lambda _, c=cat, lw=lw: self.add_files(c, lw))
            btn_remove = QPushButton("Remove Selected")
            btn_remove.clicked.connect(lambda _, lw=lw, c=cat: self.remove_selected(lw, c))
            btns.addWidget(btn_add)
            btns.addWidget(btn_remove)
            btns.addStretch()

            row.addWidget(label)
            row.addWidget(lw)
            row.addLayout(btns)
            layout.addLayout(row)
            self.lists_widgets[cat] = lw

        action_row = QHBoxLayout()
        btn_pdf = QPushButton("Export Checklist PDF")
        btn_pdf.clicked.connect(self.export_checklist_pdf)
        btn_zip = QPushButton("Perform Handover (Create ZIP + BOM + Project PDF)")
        btn_zip.clicked.connect(self.perform_handover)
        action_row.addWidget(btn_pdf)
        action_row.addWidget(btn_zip)
        action_row.addStretch()
        layout.addLayout(action_row)

        self.logo_label = QLabel()
        self.logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                if self.logo_path.lower().endswith(".gif"):
                    self.movie = QMovie(self.logo_path)
                    self.logo_label.setMovie(self.movie)
                else:
                    pix = QPixmap(self.logo_path).scaledToHeight(64, Qt.TransformationMode.SmoothTransformation)
                    self.logo_label.setPixmap(pix)
            except Exception:
                pass
        layout.addWidget(self.logo_label)
        self.setLayout(layout)

    def load_docs_for_project(self, project_id: int):
        # clear all and load from db
        for cat, lw in self.lists_widgets.items():
            lw.clear()
        docs = self.db.get_handover_docs(project_id)
        for d in docs:
            cat = d["category"]
            path = d["file_path"]
            if cat in self.lists_widgets:
                self.lists_widgets[cat].addItem(f"{d['id']}::{path}")

    def add_files(self, category: str, list_widget: QListWidget):
        try:
            files, _ = QFileDialog.getOpenFileNames(self, f"Select files for {category}")
            if not files:
                return

            proj_dir = self.get_project_dir()
            if not proj_dir:
                QMessageBox.warning(self, "No Project Selected", "Please select and confirm product/project in Project tab first.")
                return

            os.makedirs(os.path.join(proj_dir, category.replace(" ", "_")), exist_ok=True)

            progress = QProgressDialog(f"Copying {len(files)} file(s) to {category}...", None, 0, len(files), self)
            progress.setWindowModality(Qt.WindowModality.ApplicationModal)
            progress.show()

            # ✅ use project getter and db directly
            # if your get_project_info() returns (product, project, ...):
            info = self.get_project_info()
            project_name = info[1] if info else None   # index 1 = project name
            proj_row = self.db.get_project_by_name(project_name) if project_name else None
            project_id = proj_row["project_id"] if proj_row else None

            for i, f in enumerate(files, start=1):
                try:
                    fname = os.path.basename(f)
                    dest = os.path.join(proj_dir, category.replace(" ", "_"), fname)
                    shutil.copy2(f, dest)

                    rel = os.path.relpath(dest, proj_dir)
                    list_widget.addItem(rel)

                    if project_id:
                        self.db.add_handover_doc(project_id, category, rel)

                except Exception as e:
                    QMessageBox.critical(self, "Copy Error", f"Failed to copy {f}\n{e}")

                progress.setValue(i)
                QApplication.processEvents()

            progress.close()
            QMessageBox.information(self, "Files Added", f"{len(files)} file(s) added to {category}.")
        except Exception as e:
            QMessageBox.warning(self, "Open Error", f"Could not open file:\n{e}")


    def remove_selected(self, list_widget: QListWidget, category: str):
        selected = list_widget.selectedItems()
        if not selected:
            return
        proj_dir = self.get_project_dir()
        project_name = os.path.basename(proj_dir) if proj_dir else None
        proj_row = self.parent().db.get_project_by_name(project_name) if project_name else None
        project_id = proj_row["project_id"] if proj_row else None
        removed = 0
        for item in selected:
            text = item.text()
            # if stored as "id::path" try to parse
            doc_id = None
            if "::" in text:
                try:
                    doc_id = int(text.split("::", 1)[0])
                except Exception:
                    doc_id = None
            # remove from UI
            row = list_widget.row(item)
            list_widget.takeItem(row)
            # remove from DB
            if doc_id:
                try:
                    self.db.remove_handover_doc(doc_id)
                    removed += 1
                except Exception:
                    pass
            else:
                # if we have project_id, remove matching record
                if project_id:
                    # try to remove by file_path
                    file_path = text
                    cur = self.db.conn.cursor()
                    cur.execute("DELETE FROM handover_docs WHERE project_id=? AND file_path=?", (project_id, file_path))
                    self.db.conn.commit()
                    removed += 1
        QMessageBox.information(self, "Removed", f"{removed} file(s) removed from {category}.")

    def export_checklist_pdf(self):
        proj_dir = self.get_project_dir()
        if not proj_dir:
            QMessageBox.warning(self, "No Project Selected", "Select and confirm product/project first (Project tab).")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Handover Checklist", f"Handover_Checklist_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        progress = QProgressDialog("Preparing Checklist PDF...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModality.ApplicationModal)
        progress.show()
        try:
            doc = SimpleDocTemplate(save_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elems = []
            elems.append(Paragraph(f"Handover Checklist - {os.path.basename(proj_dir)}", styles["Title"]))
            elems.append(Spacer(1, 8))
            elems.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))
            elems.append(Spacer(1, 12))
            data = [["Category", "Files (paths)"]]
            for cat, lw in self.lists_widgets.items():
                files = [lw.item(i).text() for i in range(lw.count())]
                data.append([cat, "\n".join(files) if files else "No files"])
            table = Table(data, colWidths=[150, 350])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F4F4F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.black),
            ]))
            elems.append(table)
            doc.build(elems)
            QMessageBox.information(self, "Exported", f"Checklist PDF saved to:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "PDF Error", f"Failed export:\n{e}\n{traceback.format_exc()}")
        finally:
            progress.close()

    def perform_handover(self):
        proj_dir = self.get_project_dir()
        if not proj_dir:
            QMessageBox.warning(self, "No Project Selected", "Select and confirm product/project first (Project tab).")
            return
        # pick BOM from existing Excel (if exists)
        excel_path = EXCEL_FILE if os.path.exists(EXCEL_FILE) else None
        chosen_bom = None
        if excel_path:
            try:
                xls = pd.ExcelFile(excel_path)
                sheet_names = [s for s in xls.sheet_names if s != PRODUCT_SHEET]
                if sheet_names:
                    sheet, ok = QInputDialog.getItem(self, "Select BOM sheet (optional)", "Choose BOM sheet to include (or Cancel to skip):", sheet_names, 0, False)
                    if ok and sheet:
                        chosen_bom = sheet
            except Exception:
                chosen_bom = None

        progress = QProgressDialog("Performing handover tasks...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModality.ApplicationModal)
        progress.show()
        try:
            if chosen_bom:
                try:
                    df_bom = pd.read_excel(excel_path, sheet_name=chosen_bom)
                    bom_export_path = os.path.join(proj_dir, f"{chosen_bom}.xlsx")
                    with pd.ExcelWriter(bom_export_path, engine="openpyxl", mode="w") as writer:
                        df_bom.to_excel(writer, sheet_name=chosen_bom, index=False)
                except Exception as e:
                    QMessageBox.warning(self, "BOM Export Warning", f"Could not export BOM sheet '{chosen_bom}':\n{e}")
            pdf_path = os.path.join(proj_dir, "Project_Report.pdf")
            success = self._generate_project_pdf(pdf_path)
            if not success:
                QMessageBox.warning(self, "PDF Warning", "Project PDF generation had issues (see console).")
        except Exception as e:
            QMessageBox.critical(self, "Handover Error", f"Failed preparing BOM/PDF:\n{e}\n{traceback.format_exc()}")
            progress.close()
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Handover ZIP", f"Handover_{os.path.basename(proj_dir)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip", "ZIP Files (*.zip)")
        if not save_path:
            progress.close()
            return
        try:
            with zipfile.ZipFile(save_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, _, files in os.walk(proj_dir):
                    for f in files:
                        full = os.path.join(root, f)
                        arc = os.path.relpath(full, os.path.dirname(proj_dir))
                        zf.write(full, arc)
            QMessageBox.information(self, "Handover Created", f"Handover ZIP saved to:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "ZIP Error", f"Failed to create ZIP:\n{e}\n{traceback.format_exc()}")
        finally:
            progress.close()

    def _generate_project_pdf(self, pdf_path: str) -> bool:
        try:
            prod, proj, details_dict, mes_dict, build_matrix, machine_matrix = self.get_project_info()
            doc = SimpleDocTemplate(pdf_path, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=70, bottomMargin=40)
            styles = getSampleStyleSheet()
            normal_style = styles["Normal"]
            wrap_style = ParagraphStyle("wrap", parent=normal_style, fontSize=9, leading=11)
            elems = []
            elems.append(Paragraph(f"<b>Project Report - {prod} / {proj}</b>", styles["Title"]))
            elems.append(Spacer(1, 6))
            elems.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
            elems.append(Spacer(1, 12))

            elems.append(Paragraph("Project Details", styles["Heading2"]))
            proj_data = [["Field", "Value"]]
            for k in ["Project Name", "FG Part Number", "PCBA Part Number", "Start Date", "End Date", "BOM File", "NPI Engineer"]:
                proj_data.append([k, Paragraph(str(details_dict.get(k, "")), wrap_style)])
            t = Table(proj_data, colWidths=[150, 350])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a3b3d")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("ALIGN", (0, 0), (-1, 0), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("FONTSIZE", (0, 0), (-1, -1), 9), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])]))
            elems.append(t); elems.append(Spacer(1, 12))

            elems.append(Paragraph("MES Workflow Details", styles["Heading2"]))
            mes_data = [["Field", "Value"]]
            for k in ["LOT ID", "Workflow SMT - Name", "Workflow TLA - Name", "SMT - Work Order", "TLA - Work Order", "Work Order Quantity", "PO NUMBER", "PO Quantity"]:
                mes_data.append([k, Paragraph(str(mes_dict.get(k, "")), wrap_style)])
            t = Table(mes_data, colWidths=[150, 350])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a3b3d")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])]))
            elems.append(t); elems.append(Spacer(1, 12))

            elems.append(Paragraph("Build Matrix", styles["Heading2"]))
            bm_data = [["No.", "Component", "Make"]]
            for idx, row in enumerate(build_matrix, start=1):
                bm_data.append([str(idx), Paragraph(row[0] or "", wrap_style), Paragraph(row[1] or "", wrap_style)])
            t = Table(bm_data, colWidths=[40, 260, 200])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a3b3d")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])]))
            elems.append(t); elems.append(Spacer(1, 12))

            elems.append(Paragraph("Machine Programs", styles["Heading2"]))
            mm_data = [["No.", "Machine Name", "Program Name"]]
            for idx, row in enumerate(machine_matrix, start=1):
                mm_data.append([str(idx), Paragraph(row[0] or "", wrap_style), Paragraph(row[1] or "", wrap_style)])
            t = Table(mm_data, colWidths=[40, 260, 200])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a3b3d")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])]))
            elems.append(t); elems.append(Spacer(1, 12))

            elems.append(Paragraph("Handover Documents Summary", styles["Heading2"]))
            data = [["Category", "Files"]]
            # gather docs from DB for current project
            proj_dir = self.get_project_dir()
            project_name = os.path.basename(proj_dir) if proj_dir else None
            proj_row = self.parent().db.get_project_by_name(project_name) if project_name else None
            project_id = proj_row["project_id"] if proj_row else None
            if project_id:
                docs = self.db.get_handover_docs(project_id)
                grouped: Dict[str, List[str]] = {}
                for d in docs:
                    grouped.setdefault(d["category"], []).append(d["file_path"])
                for cat, fls in grouped.items():
                    files_text = "<br/>".join(fls) if fls else "No files"
                    data.append([cat, Paragraph(files_text, wrap_style)])
            else:
                for cat, lw in self.lists_widgets.items():
                    files = [lw.item(i).text() for i in range(lw.count())]
                    files_text = "<br/>".join(files) if files else "No files"
                    data.append([cat, Paragraph(files_text, wrap_style)])

            t = Table(data, colWidths=[150, 350])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a3b3d")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey])]))
            elems.append(t)

            logo_path = os.path.join(os.getcwd(), "tsat.png")

            def add_logo(canvas, doc):
                if os.path.exists(logo_path):
                    canvas.drawImage(logo_path, A4[0] - 120, A4[1] - 60, width=80, height=30, preserveAspectRatio=True, mask="auto")

            doc.build(elems, onFirstPage=add_logo, onLaterPages=add_logo)
            return True
        except Exception as e:
            print("PDF error:", e, traceback.format_exc())
            return False


# ----------------- BOM Viewer -----------------
class BOMViewer(QMainWindow):
    def __init__(self, bom_df: pd.DataFrame, parent_main_window=None):
        super().__init__(parent_main_window)
        self.bom_df = bom_df.fillna("")
        self.setWindowTitle("BOM Viewer")
        self.resize(900, 600)
        self._build_ui()

    def _build_ui(self):
        w = QWidget()
        layout = QVBoxLayout()
        header = QLabel("Bill of Materials")
        header.setObjectName("HeaderLabel")
        layout.addWidget(header)

        search_row = QHBoxLayout()
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Search in BOM...")
        self.search_box.textChanged.connect(self.filter_items)
        search_row.addWidget(QLabel("Search:"))
        search_row.addWidget(self.search_box)
        layout.addLayout(search_row)

        self.tree = QTreeWidget()
        cols = list(self.bom_df.columns)
        self.tree.setColumnCount(len(cols))
        self.tree.setHeaderLabels(cols)
        self.tree.header().sectionClicked.connect(self.sort_by_column)
        layout.addWidget(self.tree)
        self._load_data()

        btn_row = QHBoxLayout()
        btn_pdf = QPushButton("Export to PDF")
        btn_pdf.clicked.connect(self.export_to_pdf)
        btn_docx = QPushButton("Export to DOCX")
        btn_docx.clicked.connect(self.export_to_docx)
        btn_row.addWidget(btn_pdf)
        btn_row.addWidget(btn_docx)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        w.setLayout(layout)
        self.setCentralWidget(w)

    def _load_data(self):
        self.tree.clear()
        for _, row in self.bom_df.iterrows():
            row_data = [str(row[col]) for col in self.bom_df.columns]
            it = QTreeWidgetItem(row_data)
            self.tree.addTopLevelItem(it)
        self.tree.expandAll()

    def filter_items(self):
        text = self.search_box.text().lower()
        for i in range(self.tree.topLevelItemCount()):
            it = self.tree.topLevelItem(i)
            visible = False
            for c in range(self.tree.columnCount()):
                if text in it.text(c).lower():
                    visible = True
                    break
            it.setHidden(not visible)

    def sort_by_column(self, col):
        order = self.tree.header().sortIndicatorOrder()
        new_order = Qt.SortOrder.DescendingOrder if order == Qt.SortOrder.AscendingOrder else Qt.SortOrder.AscendingOrder
        self.tree.sortItems(col, new_order)

    def export_to_docx(self):
        save_path, _ = QFileDialog.getSaveFileName(self, "Save DOCX", f"BOM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", "Word Files (*.docx)")
        if not save_path:
            return
        try:
            doc = Document()
            doc.add_heading("Bill of Materials", level=1)
            table = doc.add_table(rows=1, cols=len(self.bom_df.columns))
            hdr = table.rows[0].cells
            for i, c in enumerate(self.bom_df.columns):
                hdr[i].text = str(c)
            for _, row in self.bom_df.iterrows():
                r = table.add_row().cells
                for i, c in enumerate(self.bom_df.columns):
                    r[i].text = str(row[c])
            doc.save(save_path)
            QMessageBox.information(self, "Saved", f"DOCX saved to:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "DOCX Error", f"Failed to export DOCX:\n{e}")

    def export_to_pdf(self):
        save_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", f"BOM_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", "PDF Files (*.pdf')")
        if not save_path:
            return
        progress = QProgressDialog("Exporting PDF...", None, 0, 0, self)
        progress.setWindowModality(Qt.WindowModality.ApplicationModal)
        progress.show()
        try:
            doc = SimpleDocTemplate(save_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elems = [Paragraph("Bill of Materials", styles["Title"]), Spacer(1, 8)]
            data = [list(self.bom_df.columns)]
            for _, row in self.bom_df.iterrows():
                data.append([str(row[c]) for c in self.bom_df.columns])
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F4F4F")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("ALIGN", (0, 0), (-1, -1), "LEFT"), ("GRID", (0, 0), (-1, -1), 0.25, colors.black),]))
            elems.append(table)
            doc.build(elems)
            QMessageBox.information(self, "Saved", f"PDF saved to:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "PDF Error", f"Failed to export PDF:\n{e}")
        finally:
            progress.close()


# ----------------- Checklist Tab -----------------
class ChecklistTab(QWidget):
    def __init__(self, db: DBManager, get_project_dir_callable, get_current_project_callable, parent=None):
        super().__init__(parent)
        self.db = db
        self.get_project_dir = get_project_dir_callable
        self.get_current_project = get_current_project_callable  # 👈 new
        self.tree = None


        # Template mapping (same as you had)
        self.template = {
            "Design Record (BOM & 3D/2D Drawings)": {"completed": False, "person": "SANTHOSH", "reference": ""},
            "Engineering Change Notice": {"completed": False, "person": "ADITYA", "reference": ""},
            "Customer Engineering Approval": {"completed": False, "person": "TPM", "reference": ""},
            "Process Flow Diagram": {"completed": False, "person": "SIVA/ JHON", "reference": ""},
            "Process FMEA": {"completed": False, "person": "SIVA/ JHON", "reference": ""},
            "Control Plan": {"completed": False, "person": "SIVA/ JHON", "reference": ""},
            "MSA Plan and report": {"completed": False, "person": "SIVA/ JHON", "reference": ""},
            "Dimensional Results": {"completed": False, "person": "SANTHOSH", "reference": ""},
            "Material / Performance Result": {"completed": False, "person": "SUPPLIER - CONFORMANCE REPORT", "reference": ""},
            "Initial Process Study plan and report": {"completed": False, "person": "", "reference": ""},
            "Qualified Laboratory Documentation": {"completed": False, "person": "", "reference": ""},
            "Appearance Approval Report (IF APPLICABLE)": {"completed": False, "person": "", "reference": ""},
            "Sample Production Parts": {"completed": False, "person": "SANTHOSH", "reference": ""},
            "Record of Compliance": {"completed": False, "person": "SILAMBARASAN", "reference": ""},
            "Part Submission Warrant": {"completed": False, "person": "", "reference": ""},
            "IMDS Data": {"completed": False, "person": "SILAMBARASAN", "reference": ""},
            "Packaging Requirements": {"completed": False, "person": "VIGNESH/ PREM", "reference": ""},
            "Warranty period sign off": {"completed": False, "person": "", "reference": ""},
            "Child parts drawing in DCC": {"completed": False, "person": "", "reference": ""},
            "Process drawing in DCC": {"completed": False, "person": "", "reference": ""},
            "Child parts SIR clearance": {"completed": False, "person": "", "reference": ""},
            "Fixtures and Tooling status": {"completed": False, "person": "KRISHNA PRASAD", "reference": ""},
            "Fixtures and Tooling validation report": {"completed": False, "person": "KRISHNA PRASAD", "reference": ""},
            "Test machine, Program verification & validation report": {"completed": False, "person": "SOMASHEKAR", "reference": ""},
            "Reference Master Sample": {"completed": False, "person": "SANTHOSH", "reference": ""},
            "SOP": {"completed": False, "person": "SIVA/ JHON", "reference": ""},
            "Customer feedback report of proto samples / outline drawing submitted": {"completed": False, "person": "TPM", "reference": ""},
            "Reliability and Temperature Test - PV": {"completed": False, "person": "", "reference": ""},
            "Is the manufacturing layout been reviewed to manufacture the product ?": {"completed": False, "person": "DEEPAK", "reference": ""},
            "Has the supplier Assessment been completed": {"completed": False, "person": "SDE", "reference": ""},
            "Child part agreement of Inspection (AOI)": {"completed": False, "person": "", "reference": ""},
            "Manufacturing equipments & their control facilities": {"completed": False, "person": "", "reference": ""},
            "Special process details": {"completed": False, "person": "", "reference": ""},
            "Master list of machine PM and calibration": {"completed": False, "person": "PRABHU", "reference": ""},
            "Intimation for PPAP": {"completed": False, "person": "NPI", "reference": ""},
            "Updated PPAP time line plan": {"completed": False, "person": "TPM", "reference": ""},
            "Manufacturing feasibility": {"completed": False, "person": "", "reference": ""},
            "Run @ rate - PPAP Quantity": {"completed": False, "person": "DEEPAK", "reference": ""},
            "CSR - Customer Specific Requirement": {"completed": False, "person": "", "reference": ""},
            "Safe launch": {"completed": False, "person": "SANTHOSH", "reference": ""},
            "Supplier Master List": {"completed": False, "person": "SANTHOSH", "reference": ""},
            "Lesson Learnt": {"completed": False, "person": "SANTHOSH", "reference": ""}
        }
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout()
        header = QLabel("Project Checklist")
        header.setObjectName("HeaderLabel")
        layout.addWidget(header)

        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Completed", "Checklist Item", "Person In Charge", "Reference Path"])
        self.tree.itemDoubleClicked.connect(self._handle_double_click)
        layout.addWidget(self.tree)

        header = self.tree.header()
        header.setSectionResizeMode(0, header.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(1, header.ResizeMode.Stretch)
        header.setSectionResizeMode(2, header.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(3, header.ResizeMode.Stretch)

        btn_row = QHBoxLayout()
        btn_add = QPushButton("Add Reference")
        btn_add.clicked.connect(self._add_reference)
        btn_edit = QPushButton("Edit Person")
        btn_edit.clicked.connect(self._edit_person)
        btn_save = QPushButton("Save Checklist")
        btn_save.clicked.connect(self._save_to_db)
        btn_export = QPushButton("Export Checklist PDF")
        btn_export.clicked.connect(self._export_pdf)
        btn_row.addWidget(btn_add)
        btn_row.addWidget(btn_edit)
        btn_row.addWidget(btn_save)
        btn_row.addWidget(btn_export)
        btn_row.addStretch()
        layout.addLayout(btn_row)
        self.setLayout(layout)

    def load_for_project(self, project_id: int):
        # initialize checklist if needed
        self.db.initialize_checklist(project_id, self.template)
        self._load_from_db(project_id)

    def _load_from_db(self, project_id: int):
        self.tree.clear()
        rows = self.db.get_checklist(project_id)
        for r in rows:
            it = QTreeWidgetItem(["✔" if r["completed"] else "", r["item_name"], r["person"] or "", r["reference"] or ""])
            it.setData(0, Qt.ItemDataRole.UserRole, r["id"])
            it.setCheckState(0, Qt.CheckState.Checked if r["completed"] else Qt.CheckState.Unchecked)
            self.tree.addTopLevelItem(it)
    def _save_to_db(self):
        try:    
            proj_dir = self.get_project_dir()
            if not proj_dir:
                QMessageBox.warning(self, "No Project", "Select and confirm a project first.")
                return

            project_name = self.get_current_project()  # 👈 safe getter
            if not project_name:
                QMessageBox.critical(self, "Missing", "No project selected.")
                return

            proj_row = self.db.get_project_by_name(project_name)
            if not proj_row:
                QMessageBox.critical(self, "Missing", f"Project '{project_name}' not found in DB.")
                return

            pid = proj_row["project_id"]

            for i in range(self.tree.topLevelItemCount()):
                it = self.tree.topLevelItem(i)
                item_id = it.data(0, Qt.ItemDataRole.UserRole)
                completed = 1 if it.checkState(0) == Qt.CheckState.Checked else 0
                person = it.text(2).strip()
                reference = it.text(3).strip()

                if item_id:
                    self.db.update_checklist_item(item_id, completed, person, reference)
                else:
                    new_id = self.db.insert_checklist_item(
                        pid, it.text(1), completed, person, reference
                    )
                    it.setData(0, Qt.ItemDataRole.UserRole, new_id)

            QMessageBox.information(self, "Saved", "Checklist saved successfully.")
        except Exception as e:
            QMessageBox.warning(self, "Open Error", f"Could not open file:\n{e}")


    def _add_reference(self):
        it = self.tree.currentItem()
        if not it:
            return
        fpath, _ = QFileDialog.getOpenFileName(self, "Select reference file")
        if not fpath:
            return
        it.setText(3, fpath)

    def _edit_person(self):
        it = self.tree.currentItem()
        if not it:
            return
        new_person, ok = QInputDialog.getText(self, "Edit Person", "Enter new person name:", text=it.text(2))
        if ok and new_person.strip():
            it.setText(2, new_person.strip())

    def _handle_double_click(self, item: QTreeWidgetItem, col: int):
        if col == 3 and item.text(3):
            try:
                path = item.text(3)
                if sys.platform.startswith("win"):
                    os.startfile(path)
                elif sys.platform.startswith("darwin"):
                    os.system(f'open "{path}"')
                else:
                    os.system(f'xdg-open "{path}"')
            except Exception as e:
                QMessageBox.warning(self, "Open Error", f"Could not open file:\n{e}")

    def _export_pdf(self):
        proj_dir = self.get_project_dir()
        project_name = os.path.basename(proj_dir) if proj_dir else "Unknown"
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Checklist PDF", f"Checklist_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            doc = SimpleDocTemplate(save_path, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=70, bottomMargin=50)
            styles = getSampleStyleSheet()
            normal = styles["Normal"]
            wrap = ParagraphStyle("wrap", parent=normal, fontSize=9, leading=11)
            elems = []
            elems.append(Paragraph(f"<b>Project Checklist - {project_name}</b>", styles["Title"]))
            elems.append(Spacer(1, 12))
            data = [["S.No", "Completed", "Checklist Item", "Person", "Reference"]]
            for i in range(self.tree.topLevelItemCount()):
                it = self.tree.topLevelItem(i)
                data.append([str(i + 1), "✔" if it.checkState(0) == Qt.CheckState.Checked else "", Paragraph(it.text(1), wrap), Paragraph(it.text(2), wrap), "Available" if it.text(3).strip() else ""])
            table = Table(data, repeatRows=1, colWidths=[40, 60, 220, 120, 100])
            table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F4F4F")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.25, colors.black), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("ALIGN", (0, 0), (1, -1), "CENTER"),]))
            elems.append(table)
            logo_path = os.path.join(os.getcwd(), "tsat.png")
            def add_decorations(canvas, doc):
                canvas.setFont("Helvetica", 8)
                footer_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Created by VVS"
                canvas.drawString(40, 30, footer_text)
                if os.path.exists(logo_path):
                    canvas.drawImage(logo_path, A4[0] - 140, A4[1] - 80, width=100, height=50, preserveAspectRatio=True, mask="auto")
            doc.build(elems, onFirstPage=add_decorations, onLaterPages=add_decorations)
            QMessageBox.information(self, "Exported", f"Checklist PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "PDF Error", f"Failed to export PDF:\n{e}")


# ----------------- Main Application -----------------
class NPIProjectManager(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("NPI Project Manager (DB)")
        self.resize(1080, 720)
        self.db = DBManager(DB_FILE)
        self.current_product = None
        self.current_project = None

        self._build_ui()
        self.load_products()

    def _build_ui(self):
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        proj_tab = QWidget()
        proj_layout = QVBoxLayout()
        proj_tab.setLayout(proj_layout)
        header = QLabel("Project Manager")
        header.setObjectName("HeaderLabel")
        proj_layout.addWidget(header)

        sel_row = QHBoxLayout()
        sel_row.addWidget(QLabel("Product:"))
        self.product_dropdown = QComboBox()
        self.product_dropdown.setMinimumWidth(250)
        sel_row.addWidget(self.product_dropdown)
        self.btn_add_product = QPushButton("Add Product")
        self.btn_add_product.clicked.connect(self.add_product)
        sel_row.addWidget(self.btn_add_product)
        self.btn_confirm_product = QPushButton("Confirm Product")
        self.btn_confirm_product.clicked.connect(self.confirm_product)
        sel_row.addWidget(self.btn_confirm_product)

        sel_row.addSpacing(20)
        sel_row.addWidget(QLabel("Project:"))
        self.project_dropdown = QComboBox()
        self.project_dropdown.setMinimumWidth(300)
        sel_row.addWidget(self.project_dropdown)
        self.btn_view_project = QPushButton("View Project")
        self.btn_view_project.clicked.connect(self.view_project)
        sel_row.addWidget(self.btn_view_project)
        self.btn_add_project = QPushButton("Add Project")
        self.btn_add_project.clicked.connect(self.add_project)
        sel_row.addWidget(self.btn_add_project)
        proj_layout.addLayout(sel_row)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        left = QWidget()
        left_layout = QVBoxLayout()
        left.setLayout(left_layout)
        left_layout.addWidget(QLabel("<b>Project Details</b>"))
        self.details_fields = ["Project Name", "FG Part Number", "PCBA Part Number", "Start Date", "End Date", "BOM File", "NPI Engineer"]
        self.details_entries: Dict[str, QLineEdit] = {}
        for f in self.details_fields:
            row = QHBoxLayout()
            row.addWidget(QLabel(f))
            le = QLineEdit()
            le.setReadOnly(True)
            le.setFixedWidth(320)
            self.details_entries[f] = le
            row.addWidget(le)
            left_layout.addLayout(row)
        self.btn_edit_project = QPushButton("Edit Project")
        self.btn_edit_project.clicked.connect(self.enable_editing)
        left_layout.addWidget(self.btn_edit_project)
        self.btn_update_project = QPushButton("Update Project")
        self.btn_update_project.clicked.connect(self.update_project)
        self.btn_update_project.setEnabled(False)
        left_layout.addWidget(self.btn_update_project)
        splitter.addWidget(left)

        right = QWidget()
        right_layout = QVBoxLayout()
        right.setLayout(right_layout)
        right_layout.addWidget(QLabel("<b>MES Workflow Details</b>"))
        self.mes_fields = ["LOT ID", "Workflow SMT - Name", "Workflow TLA - Name", "SMT - Work Order", "TLA - Work Order", "Work Order Quantity", "PO NUMBER", "PO Quantity"]
        self.mes_entries: Dict[str, QLineEdit] = {}
        for f in self.mes_fields:
            row = QHBoxLayout()
            row.addWidget(QLabel(f))
            le = QLineEdit()
            le.setReadOnly(True)
            le.setFixedWidth(320)
            self.mes_entries[f] = le
            row.addWidget(le)
            right_layout.addLayout(row)
        splitter.addWidget(right)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 1)
        proj_layout.addWidget(splitter)

        matrix_row = QHBoxLayout()
        assembly_box = QVBoxLayout()
        assembly_box.addWidget(QLabel("<b>Assembly Drawings</b>"))
        self.assembly_table = QTableWidget(ASSEMBLY_ROWS, 2)
        self.assembly_table.setHorizontalHeaderLabels(["Assembly Drawing", "Drawing Name"])
        try:
            from PyQt6.QtWidgets import QHeaderView
            self.assembly_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        except Exception:
            pass
        matrix_row.addLayout(assembly_box, 2)

        build_box = QVBoxLayout()
        build_box.addWidget(QLabel("<b>Build Matrix</b>"))
        self.build_matrix_table = QTableWidget(TABLE_ROWS, 2)
        self.build_matrix_table.setHorizontalHeaderLabels(["Component", "Make"])
        try:
            from PyQt6.QtWidgets import QHeaderView
            self.build_matrix_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        except Exception:
            pass
        build_box.addWidget(self.build_matrix_table)
        matrix_row.addLayout(build_box, 3)

        machine_box = QVBoxLayout()
        machine_box.addWidget(QLabel("<b>Machine Program</b>"))
        self.machine_program_table = QTableWidget(MACHINE_ROWS, 2)
        self.machine_program_table.setHorizontalHeaderLabels(["Machine Name", "Program Name"])
        try:
            from PyQt6.QtWidgets import QHeaderView
            self.machine_program_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        except Exception:
            pass
        machine_box.addWidget(self.machine_program_table)
        matrix_row.addLayout(machine_box, 3)
        proj_layout.addLayout(matrix_row)

        bom_row = QHBoxLayout()
        bom_row.addWidget(QLabel("<b>BOM Upload & Viewer</b>"))
        self.btn_upload_bom = QPushButton("Upload BOM")
        self.btn_upload_bom.clicked.connect(self.upload_bom)
        bom_row.addWidget(self.btn_upload_bom)
        self.btn_view_bom = QPushButton("View BOM")
        self.btn_view_bom.clicked.connect(self.view_bom)
        bom_row.addWidget(self.btn_view_bom)
        bom_row.addStretch()
        sig = QLabel("Created by - VVS")
        sig.setStyleSheet("color: #6a6a6a;")
        bom_row.addWidget(sig)
        proj_layout.addLayout(bom_row)

        self.tabs.addTab(proj_tab, "Project")

        # Handover & Checklist tabs (pass DB)
        self.handover_tab = HandoverTab(self.db, self.get_current_project_dir, self._collect_project_info, logo_path="lg.png")
        self.tabs.addTab(self.handover_tab, "Handover")
        self.checklist_tab = ChecklistTab(
            db=self.db,
            get_project_dir_callable=self.get_current_project_dir,
            get_current_project_callable=lambda: self.current_project  # 👈 pass in getter
        )

        self.tabs.addTab(self.checklist_tab, "Checklist")

        self.btn_add_assembly = QPushButton("Add Assembly Drawing(s)")
        self.btn_add_assembly.clicked.connect(self.add_assembly_drawings)
        assembly_box.addWidget(self.assembly_table)
        assembly_box.addWidget(self.btn_add_assembly)

        self.setStyleSheet(PROFESSIONAL_QSS)

    # Helper
    def get_current_project_dir(self):
        if not self.current_product or not self.current_project:
            return None
        safe = f"{self.current_product}_{self.current_project}".replace(" ", "_")
        proj_dir = os.path.join(os.getcwd(), "Projects", safe)
        os.makedirs(proj_dir, exist_ok=True)
        return proj_dir

    # Load products from DB (fall back to Excel if DB empty)
    def load_products(self):
        try:
            products = self.db.list_products()
            if not products and os.path.exists(EXCEL_FILE):
                xls = pd.ExcelFile(EXCEL_FILE)
                sheets = xls.sheet_names
                if PRODUCT_SHEET in sheets:
                    try:
                        df = pd.read_excel(EXCEL_FILE, sheet_name=PRODUCT_SHEET)
                        products = df["Product Name"].dropna().astype(str).tolist()
                    except Exception:
                        products = sheets
                else:
                    products = sheets
            self.product_dropdown.clear()
            self.product_dropdown.addItems(products)
        except Exception as e:
            QMessageBox.critical(self, "Load Error", f"Failed to load products:\n{e}")

    def confirm_product(self):
        self.current_product = self.product_dropdown.currentText()
        if not self.current_product:
            QMessageBox.warning(self, "Select Product", "Please select a product.")
            return
        self.load_projects_for_product()

    def load_projects_for_product(self):
        try:
            projects = self.db.list_projects_for_product(self.current_product)
            # fallback to excel sheet
            if not projects and os.path.exists(EXCEL_FILE):
                df = pd.read_excel(EXCEL_FILE, sheet_name=self.current_product)
                if "Project Name" in df.columns:
                    projects = df["Project Name"].dropna().unique().tolist()
            self.project_dropdown.clear()
            self.project_dropdown.addItems(projects)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load projects:\n{e}")

    def add_product(self):
        prod, ok = QInputDialog.getText(self, "Add Product", "Enter new product name:")
        if not ok or not prod.strip():
            return
        prod = prod.strip()
        # create a dummy project row or no-op: we will just add to dropdown and DB when project added
        self.product_dropdown.addItem(prod)
        QMessageBox.information(self, "Added", f"Product '{prod}' added to dropdown. Add a project to persist to DB.")

    def add_project(self):
        if not self.current_product:
            QMessageBox.warning(self, "Select Product", "Select a product first.")
            return
        proj, ok = QInputDialog.getText(self, "New Project", "Enter new project name:")
        if not ok or not proj.strip():
            return
        proj = proj.strip()
        # collect details from UI fields (they may be empty)
        row = {f: self.details_entries[f].text() for f in self.details_fields}
        row.update({f: self.mes_entries[f].text() for f in self.mes_fields})
        # collect tables
        build_rows = []
        for i in range(TABLE_ROWS):
            c_item = self.build_matrix_table.item(i, 0)
            m_item = self.build_matrix_table.item(i, 1)
            build_rows.append((c_item.text() if c_item else "", m_item.text() if m_item else ""))
        assembly_rows = []
        for i in range(ASSEMBLY_ROWS):
            a = self.assembly_table.item(i, 0)
            n = self.assembly_table.item(i, 1)
            assembly_rows.append((a.text() if a else "", n.text() if n else ""))
        machine_rows = []
        for i in range(MACHINE_ROWS):
            mn = self.machine_program_table.item(i, 0)
            pn = self.machine_program_table.item(i, 1)
            machine_rows.append((mn.text() if mn else "", pn.text() if pn else ""))
        try:
            pid = self.db.add_project(self.current_product, proj, row)
            # save MES, build, machine
            self.db.save_mes(pid, {k: self.mes_entries[k].text() for k in self.mes_fields})
            self.db.save_assembly_drawings(pid, assembly_rows)
            self.db.save_build_matrix(pid, build_rows)
            self.db.save_machine_matrix(pid, machine_rows)
            # initialize checklist
            self.db.initialize_checklist(pid, self.checklist_tab.template)
            QMessageBox.information(self, "Added", "Project added to DB.")
            self.load_projects_for_product()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add project:\n{e}\n{traceback.format_exc()}")

    def view_project(self):
        self.current_project = self.project_dropdown.currentText()
        self.current_product = self.product_dropdown.currentText()
        if not self.current_product or not self.current_project:
            QMessageBox.warning(self, "Select", "Please select both product and project.")
            return
        try:
            # load project from DB; if not exist, try excel fallback
            proj_row = self.db.get_project_by_name(self.current_project)
            if proj_row:
                # fill details
                for f in self.details_fields:
                    key = f.lower().replace(" ", "_")
                    val = proj_row[key] if key in proj_row.keys() else None
                    # you can use val if needed later

                    # fallback mapping:
                # direct set using columns we know
                self.details_entries["Project Name"].setText(str(proj_row["project_name"] or ""))
                self.details_entries["FG Part Number"].setText(str(proj_row["fg_part_number"] or ""))
                self.details_entries["PCBA Part Number"].setText(str(proj_row["pcba_part_number"] or ""))
                self.details_entries["Start Date"].setText(str(proj_row["start_date"] or ""))
                self.details_entries["End Date"].setText(str(proj_row["end_date"] or ""))
                self.details_entries["BOM File"].setText(str(proj_row["bom_file"] or ""))
                self.details_entries["NPI Engineer"].setText(str(proj_row["npi_engineer"] or ""))

                mes_row = self.db.get_mes(proj_row["project_id"])
                if mes_row:
                    self.mes_entries["LOT ID"].setText(str(mes_row["lot_id"] or ""))
                    self.mes_entries["Workflow SMT - Name"].setText(str(mes_row["workflow_smt"] or ""))
                    self.mes_entries["Workflow TLA - Name"].setText(str(mes_row["workflow_tla"] or ""))
                    self.mes_entries["SMT - Work Order"].setText(str(mes_row["smt_work_order"] or ""))
                    self.mes_entries["TLA - Work Order"].setText(str(mes_row["tla_work_order"] or ""))
                    self.mes_entries["Work Order Quantity"].setText(str(mes_row["work_order_qty"] or ""))
                    self.mes_entries["PO NUMBER"].setText(str(mes_row["po_number"] or ""))
                    self.mes_entries["PO Quantity"].setText(str(mes_row["po_qty"] or ""))

                # build matrix
                for i in range(TABLE_ROWS):
                    self.build_matrix_table.setItem(i, 0, QTableWidgetItem(""))
                    self.build_matrix_table.setItem(i, 1, QTableWidgetItem(""))
                bm = self.db.get_build_matrix(proj_row["project_id"])
                for i, r in enumerate(bm):
                    self.build_matrix_table.setItem(i, 0, QTableWidgetItem(str(r["component"])))
                    self.build_matrix_table.setItem(i, 1, QTableWidgetItem(str(r["make"])))

                # assembly - we store in build_matrix? we kept assembly separate in excel; try to fetch from machine rows if any placeholder; else skip
                # assembly table
                for i in range(ASSEMBLY_ROWS):
                    self.assembly_table.setItem(i, 0, QTableWidgetItem(""))
                    self.assembly_table.setItem(i, 1, QTableWidgetItem(""))

                assembly_rows = self.db.get_assembly_drawings(proj_row["project_id"])
                for i, r in enumerate(assembly_rows):
                    if i >= ASSEMBLY_ROWS:
                        break
                    self.assembly_table.setItem(i, 0, QTableWidgetItem(str(r["assembly_drawing"])))
                    self.assembly_table.setItem(i, 1, QTableWidgetItem(str(r["drawing_name"])))

                
                # machine matrix
                for i in range(MACHINE_ROWS):
                    self.machine_program_table.setItem(i, 0, QTableWidgetItem(""))
                    self.machine_program_table.setItem(i, 1, QTableWidgetItem(""))
                mm = self.db.get_machine_matrix(proj_row["project_id"])
                for i, r in enumerate(mm):
                    self.machine_program_table.setItem(i, 0, QTableWidgetItem(str(r["machine_name"])))
                    self.machine_program_table.setItem(i, 1, QTableWidgetItem(str(r["program_name"])))

                # load handover docs and checklist
                self.handover_tab.load_docs_for_project(proj_row["project_id"])
                self.checklist_tab.load_for_project(proj_row["project_id"])
            else:
                # fallback to excel behavior (read sheet)
                df = pd.read_excel(EXCEL_FILE, sheet_name=self.current_product)
                if "Project Name" not in df.columns:
                    QMessageBox.critical(self, "Error", "Project Name column missing.")
                    return
                row = df[df["Project Name"] == self.current_project]
                if row.empty:
                    QMessageBox.warning(self, "Not Found", "Project details not found.")
                    return
                row = row.iloc[0]
                for f in self.details_fields:
                    v = row.get(f, "")
                    self.details_entries[f].setText("" if pd.isna(v) else str(v))
                for f in self.mes_fields:
                    v = row.get(f, "")
                    self.mes_entries[f].setText("" if pd.isna(v) else str(v))
                for i in range(TABLE_ROWS):
                    comp = row.get(f"Component {i+1}", "")
                    make = row.get(f"Make {i+1}", "")
                    if not pd.isna(comp) and comp != "":
                        self.build_matrix_table.setItem(i, 0, QTableWidgetItem(str(comp)))
                    if not pd.isna(make) and make != "":
                        self.build_matrix_table.setItem(i, 1, QTableWidgetItem(str(make)))
                for i in range(ASSEMBLY_ROWS):
                    ad = row.get(f"Assembly Drawing {i+1}", "")
                    dn = row.get(f"Drawing Name {i+1}", "")
                    if not pd.isna(ad) and ad != "":
                        self.assembly_table.setItem(i, 0, QTableWidgetItem(str(ad)))
                    if not pd.isna(dn) and dn != "":
                        self.assembly_table.setItem(i, 1, QTableWidgetItem(str(dn)))
                for i in range(MACHINE_ROWS):
                    mn = row.get(f"Machine Name {i+1}", "")
                    pn = row.get(f"Program Name {i+1}", "")
                    if not pd.isna(mn) and mn != "":
                        self.machine_program_table.setItem(i, 0, QTableWidgetItem(str(mn)))
                    if not pd.isna(pn) and pn != "":
                        self.machine_program_table.setItem(i, 1, QTableWidgetItem(str(pn)))
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load project:\n{e}\n{traceback.format_exc()}")

    def enable_editing(self):
        pwd, ok = QInputDialog.getText(self, "Password", "Enter the Password:", QLineEdit.EchoMode.Password)
        if not ok or pwd.strip() != "vvs18":
            QMessageBox.warning(self, "Denied", "Incorrect password.")
            return
        for e in self.details_entries.values():
            e.setReadOnly(False)
        for e in self.mes_entries.values():
            e.setReadOnly(False)
        self.btn_update_project.setEnabled(True)
        for i in range(TABLE_ROWS):
            for j in range(2):
                it = self.build_matrix_table.item(i, j)
                if it is None:
                    it = QTableWidgetItem("")
                    self.build_matrix_table.setItem(i, j, it)
                it.setFlags(it.flags() | Qt.ItemFlag.ItemIsEditable)
        for i in range(ASSEMBLY_ROWS):
            for j in range(2):
                it = self.assembly_table.item(i, j)
                if it is None:
                    it = QTableWidgetItem("")
                    self.assembly_table.setItem(i, j, it)
                it.setFlags(it.flags() | Qt.ItemFlag.ItemIsEditable)
        for i in range(MACHINE_ROWS):
            for j in range(2):
                it = self.machine_program_table.item(i, j)
                if it is None:
                    it = QTableWidgetItem("")
                    self.machine_program_table.setItem(i, j, it)
                it.setFlags(it.flags() | Qt.ItemFlag.ItemIsEditable)

    def update_project(self):
        if not self.current_product or not self.current_project:
            QMessageBox.warning(self, "Select", "Select product/project first.")
            return
        try:
            # Update DB record
            details = {f: self.details_entries[f].text() for f in self.details_fields}
            self.db.update_project_details(self.current_project, details)
            # update mes
            proj_row = self.db.get_project_by_name(self.current_project)
            if proj_row:
                pid = proj_row["project_id"]
                self.db.save_mes(pid, {k: self.mes_entries[k].text() for k in self.mes_fields})
                # save build and machine matrices
                build_rows = []
                for i in range(TABLE_ROWS):
                    ci = self.build_matrix_table.item(i, 0)
                    mi = self.build_matrix_table.item(i, 1)
                    build_rows.append((ci.text() if ci else "", mi.text() if mi else ""))
                self.db.save_build_matrix(pid, build_rows)
                machine_rows = []
                for i in range(MACHINE_ROWS):
                    mni = self.machine_program_table.item(i, 0)
                    pni = self.machine_program_table.item(i, 1)
                    machine_rows.append((mni.text() if mni else "", pni.text() if pni else ""))
                self.db.save_machine_matrix(pid, machine_rows)
                # save assembly matrix <-- NEW
                assembly_rows = []
                for i in range(ASSEMBLY_ROWS):
                    ad = self.assembly_table.item(i, 0)
                    dn = self.assembly_table.item(i, 1)
                    assembly_rows.append((ad.text() if ad else "", dn.text() if dn else ""))
                self.db.save_assembly_drawings(pid, assembly_rows)
            QMessageBox.information(self, "Saved", "Project updated.")
            for e in self.details_entries.values():
                e.setReadOnly(True)
            for e in self.mes_entries.values():
                e.setReadOnly(True)
            self.btn_update_project.setEnabled(False)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to update project:\n{e}\n{traceback.format_exc()}")

    def upload_bom(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Upload BOM", "", "Excel Files (*.xlsx *.xls)")
        if not file_path:
            return
        sheet_name, ok = QInputDialog.getText(self, "BOM Sheet Name", "Enter BOM sheet name:")
        if not ok or not sheet_name.strip():
            return
        try:
            df = pd.read_excel(file_path)
            if os.path.exists(EXCEL_FILE):
                with pd.ExcelWriter(EXCEL_FILE, mode="a", engine="openpyxl", if_sheet_exists="replace") as writer:
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            else:
                with pd.ExcelWriter(EXCEL_FILE, mode="w", engine="openpyxl") as writer:
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            QMessageBox.information(self, "Uploaded", "BOM uploaded successfully.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to upload BOM:\n{e}\n{traceback.format_exc()}")

    def view_bom(self):
        if not os.path.exists(EXCEL_FILE):
            QMessageBox.warning(self, "Missing", "No Excel datafile found.")
            return
        try:
            xls = pd.ExcelFile(EXCEL_FILE)
            sheet_names = xls.sheet_names
            excluded = {"USB DUO", "VCUSB", "GLOVE BOX", "GLOVEBOX", "test", "Dummy", "AUDIO AMPLIFIER", "BMB", PRODUCT_SHEET, "HVAC"}
            filtered = [s for s in sheet_names if s not in excluded]
            if not filtered:
                QMessageBox.warning(self, "No BOM", "No valid BOM sheets found.")
                return
            sheet, ok = QInputDialog.getItem(self, "Select BOM Sheet", "Choose BOM sheet:", filtered, 0, False)
            if not ok or not sheet:
                return
            bom_df = pd.read_excel(EXCEL_FILE, sheet_name=sheet)
            viewer = BOMViewer(bom_df, self)
            viewer.show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open BOM:\n{e}\n{traceback.format_exc()}")

    def add_assembly_drawings(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Select Assembly Drawings")
        if not files:
            return
        proj_dir = self.get_current_project_dir()
        if not proj_dir:
            QMessageBox.warning(self, "No Project", "Select and confirm a project first.")
            return
        target = os.path.join(proj_dir, "Assembly_Drawings")
        os.makedirs(target, exist_ok=True)
        progress = QProgressDialog("Copying drawings...", None, 0, len(files), self)
        progress.setWindowModality(Qt.WindowModality.ApplicationModal)
        progress.show()
        for i, f in enumerate(files, start=1):
            try:
                dest = os.path.join(target, os.path.basename(f))
                shutil.copy2(f, dest)
                placed = False
                for r in range(ASSEMBLY_ROWS):
                    it = self.assembly_table.item(r, 0)
                    if it is None or it.text() == "":
                        self.assembly_table.setItem(r, 0, QTableWidgetItem(dest))
                        self.assembly_table.setItem(r, 1, QTableWidgetItem(os.path.basename(dest)))
                        placed = True
                        break
                if not placed:
                    self.assembly_table.setItem(0, 0, QTableWidgetItem(dest))
                    self.assembly_table.setItem(0, 1, QTableWidgetItem(os.path.basename(dest)))
            except Exception as e:
                QMessageBox.critical(self, "Copy Error", f"Failed to copy {f}\n{e}")
            progress.setValue(i)
            QApplication.processEvents()
        progress.close()
        QMessageBox.information(self, "Done", f"Copied {len(files)} drawing(s).")

    def _collect_project_info(self):
        prod = self.current_product or ""
        proj = self.current_project or ""
        details = {f: (self.details_entries[f].text() if f in self.details_entries else "") for f in self.details_fields}
        mes = {f: (self.mes_entries[f].text() if f in self.mes_entries else "") for f in self.mes_fields}
        build_matrix = []
        for i in range(TABLE_ROWS):
            a = self.build_matrix_table.item(i, 0).text() if self.build_matrix_table.item(i, 0) else ""
            b = self.build_matrix_table.item(i, 1).text() if self.build_matrix_table.item(i, 1) else ""
            build_matrix.append((a, b))
        machine_matrix = []
        for i in range(MACHINE_ROWS):
            a = self.machine_program_table.item(i, 0).text() if self.machine_program_table.item(i, 0) else ""
            b = self.machine_program_table.item(i, 1).text() if self.machine_program_table.item(i, 1) else ""
            machine_matrix.append((a, b))
        return prod, proj, details, mes, build_matrix, machine_matrix

    def closeEvent(self, event):
        try:
            self.db.close()
        except Exception:
            pass
        event.accept()


def main():
    app = QApplication(sys.argv)
    if QDARK_AVAILABLE:
        try:
            app.setStyleSheet(qdarkstyle.load_stylesheet_pyqt6())
        except Exception:
            pass
    win = NPIProjectManager()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
